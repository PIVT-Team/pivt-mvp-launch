#!/usr/bin/env python3
"""Regenerate the Office fixtures used by the extraction tests.

Uses real generators (openpyxl, python-docx) on purpose. A hand-written .xlsx
only proves the parser agrees with whoever wrote the fixture.

    pip install openpyxl python-docx
    python3 scripts/make-extraction-fixtures.py
"""
import datetime, pathlib
import openpyxl
from openpyxl.styles import Font
from docx import Document

OUT = pathlib.Path(__file__).resolve().parent.parent / "supabase/functions/_shared/__tests__/fixtures"
OUT.mkdir(parents=True, exist_ok=True)

wb = openpyxl.Workbook()
sm = wb.active; sm.title = "Summary"
sm["A1"] = "Greenfield Solar Inc. — Closing Summary"; sm["A1"].font = Font(bold=True)
sm["A3"] = "Purchase Price"; sm["B3"] = 185000000; sm["B3"].number_format = "#,##0"
sm["A4"] = "Signing Date"; sm["B4"] = datetime.date(2026, 3, 14); sm["B4"].number_format = "yyyy-mm-dd"
sm["A5"] = "Escrow %"; sm["B5"] = 0.10; sm["B5"].number_format = "0.0%"

ct = wb.create_sheet("Cap Table")
ct.append(["Holder", "Class", "Shares", "Ownership %", "Payout"])
for r in [
    ("Meridian Holdings LLC", "Common", 4_200_000, 0.42, 77_700_000),
    ("Meridian Capital LLC",  "Common",   800_000, 0.08, 14_800_000),
    ("Apex Advisory, L.L.C.", "Series A", 2_500_000, 0.25, 46_250_000),
    ("Jane Okafor",           "Common",   500_000, 0.05,  9_250_000),
    ("Option Pool",           "Options", 2_000_000, 0.20, 37_000_000),
]:
    ct.append(list(r))
for i in range(2, 7):
    ct.cell(i, 4).number_format = "0.00%"
    ct.cell(i, 5).number_format = "$#,##0"
ct["A8"] = "Total"; ct["C8"] = "=SUM(C2:C6)"; ct["D8"] = 1.0; ct["D8"].number_format = "0.00%"

sp = wb.create_sheet("Sparse")
sp["A2"] = "Wire Ref"; sp["C2"] = "ACH-99120"; sp["E2"] = "Confirmed"
wb.save(OUT / "captable.xlsx")

d = Document()
d.add_heading("Stock Purchase Agreement", 0)
d.add_paragraph("This Stock Purchase Agreement (this “Agreement”) is entered into as of March 14, 2026, "
                "by and among CleanTech Ventures, Inc. (“Buyer”) and the persons listed on Schedule A (“Sellers”).")
d.add_heading("1. Purchase and Sale", level=1)
d.add_paragraph("1.1  Purchase Price. The aggregate purchase price shall be One Hundred Eighty-Five Million "
                "Dollars ($185,000,000), subject to adjustment under Section 2.4.")
d.add_paragraph("1.2  Consent. Seller shall obtain the written consent of each counterparty listed in "
                "Schedule 3.6 prior to Closing; consent of Northgate Bank is required under the Credit Agreement.")
t = d.add_table(rows=3, cols=3)
for i, row in enumerate([["Party", "Role", "Notice Email"],
                         ["CleanTech Ventures, Inc.", "Buyer", "legal@cleantech.example"],
                         ["Meridian Holdings LLC", "Seller", "ops@meridian.example"]]):
    for j, v in enumerate(row):
        t.rows[i].cells[j].text = v
d.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
d.save(OUT / "spa.docx")

n = Document()
n.add_paragraph("Schedule 3.6 — Required Consents")
outer = n.add_table(rows=2, cols=2)
outer.rows[0].cells[0].text = "Counterparty"; outer.rows[0].cells[1].text = "Detail"
outer.rows[1].cells[0].text = "Northgate Bank"
inner = outer.rows[1].cells[1].add_table(rows=2, cols=2)
inner.rows[0].cells[0].text = "Agreement"; inner.rows[0].cells[1].text = "Credit Agreement"
inner.rows[1].cells[0].text = "Section";   inner.rows[1].cells[1].text = "7.2(b)"
n.add_paragraph("End of Schedule 3.6.")
n.save(OUT / "nested.docx")

print(f"wrote fixtures to {OUT}")
